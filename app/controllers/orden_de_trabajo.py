from fastapi import APIRouter, Depends, HTTPException, Query, status
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import desc
from typing import Optional
from datetime import date
import io

from app.core.database import get_db
from app.core.auth_deps import coordinador_or_admin
from app.models.models_ordenes_trabajo import OrdenTrabajoDiario
from app.schemas.schemas_ordenes_trabajo import (
    OrdenTrabajoSchema,
    OrdenTrabajoCreate,
    OrdenTrabajoUpdate,
    OrdenTrabajoListResponse,
)

# Importaciones para generación de Word
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

router = APIRouter(
    prefix="/ordenes-trabajo",
    tags=["Órdenes de Trabajo"],
    responses={404: {"description": "Orden de trabajo no encontrada"}},
)


# ==================== RUTAS GET ====================


@router.get(
    "/",
    response_model=OrdenTrabajoListResponse,
    dependencies=[Depends(coordinador_or_admin)],
)
def get_ordenes_trabajo(
    skip: int = Query(0, ge=0, description="Número de registros a omitir"),
    limit: int = Query(
        100, ge=1, le=1000, description="Número de registros a retornar"
    ),
    fecha_inicio: Optional[date] = Query(None, description="Filtrar desde fecha"),
    fecha_fin: Optional[date] = Query(None, description="Filtrar hasta fecha"),
    db: Session = Depends(get_db),
):
    """
    Obtiene la lista de órdenes de trabajo diarias con paginación y filtros.
    Acceso restringido a Coordinadores y Administradores.
    """
    query = db.query(OrdenTrabajoDiario)

    # Aplicar filtros
    if fecha_inicio:
        query = query.filter(OrdenTrabajoDiario.fecha >= fecha_inicio)

    if fecha_fin:
        query = query.filter(OrdenTrabajoDiario.fecha <= fecha_fin)

    # Ordenar por fecha descendente
    query = query.order_by(desc(OrdenTrabajoDiario.fecha))

    # Contar total
    total = query.count()

    # Paginación
    ordenes = query.offset(skip).limit(limit).all()

    return OrdenTrabajoListResponse(
        success=True,
        total=total,
        skip=skip,
        limit=limit,
        data=ordenes,
    )


@router.get(
    "/{orden_id}",
    response_model=OrdenTrabajoSchema,
    dependencies=[Depends(coordinador_or_admin)],
)
def get_orden_by_id(
    orden_id: int,
    db: Session = Depends(get_db),
):
    """
    Obtiene una orden de trabajo específica por su ID.
    Acceso restringido a Coordinadores y Administradores.
    """
    orden = (
        db.query(OrdenTrabajoDiario).filter(OrdenTrabajoDiario.id == orden_id).first()
    )

    if not orden:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Orden de trabajo con ID {orden_id} no encontrada",
        )

    return orden


# ==================== RUTAS POST ====================


@router.post(
    "/",
    response_model=OrdenTrabajoSchema,
    status_code=status.HTTP_201_CREATED,
    dependencies=[Depends(coordinador_or_admin)],
)
def crear_orden_trabajo(
    orden: OrdenTrabajoCreate,
    db: Session = Depends(get_db),
):
    """
    Crea una nueva orden de trabajo diario.
    Acceso restringido a Coordinadores y Administradores.
    """
    # Verificar si ya existe una orden para esa fecha (opcional, pero recomendable para evitar duplicados diarios si esa es la regla)
    # En este caso permitiremos múltiples por día según requerimiento implícito de "agregar hasta 5 trabajos" que ya están en el modelo

    nueva_orden = OrdenTrabajoDiario(
        fecha=orden.fecha,
        trabajo_1_a_realizar=orden.trabajo_1_a_realizar,
        trabajo_1_realizado=orden.trabajo_1_realizado,
        trabajo_2_a_realizar=orden.trabajo_2_a_realizar,
        trabajo_2_realizado=orden.trabajo_2_realizado,
        trabajo_3_a_realizar=orden.trabajo_3_a_realizar,
        trabajo_3_realizado=orden.trabajo_3_realizado,
        trabajo_4_a_realizar=orden.trabajo_4_a_realizar,
        trabajo_4_realizado=orden.trabajo_4_realizado,
        trabajo_5_a_realizar=orden.trabajo_5_a_realizar,
        trabajo_5_realizado=orden.trabajo_5_realizado,
    )

    try:
        db.add(nueva_orden)
        db.commit()
        db.refresh(nueva_orden)
        return nueva_orden
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error al crear la orden de trabajo: {str(e)}",
        )


# ==================== RUTAS PUT ====================


@router.put(
    "/{orden_id}",
    response_model=OrdenTrabajoSchema,
    dependencies=[Depends(coordinador_or_admin)],
)
def actualizar_orden_trabajo(
    orden_id: int,
    orden_data: OrdenTrabajoUpdate,
    db: Session = Depends(get_db),
):
    """
    Actualiza una orden de trabajo existente.
    Acceso restringido a Coordinadores y Administradores.
    """
    orden = (
        db.query(OrdenTrabajoDiario).filter(OrdenTrabajoDiario.id == orden_id).first()
    )

    if not orden:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Orden de trabajo con ID {orden_id} no encontrada",
        )

    # Actualizar solo los campos proporcionados
    update_data = orden_data.model_dump(exclude_unset=True)

    for key, value in update_data.items():
        setattr(orden, key, value)

    try:
        db.commit()
        db.refresh(orden)
        return orden
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error al actualizar la orden: {str(e)}",
        )


@router.delete(
    "/{orden_id}",
    status_code=status.HTTP_200_OK,
    dependencies=[Depends(coordinador_or_admin)],
)
def eliminar_orden_trabajo(
    orden_id: int,
    db: Session = Depends(get_db),
):
    """
    Elimina una orden de trabajo.
    Acceso restringido a Coordinadores y Administradores.
    """
    orden = (
        db.query(OrdenTrabajoDiario).filter(OrdenTrabajoDiario.id == orden_id).first()
    )

    if not orden:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Orden de trabajo con ID {orden_id} no encontrada",
        )

    try:
        db.delete(orden)
        db.commit()
        return {"success": True, "message": "Orden eliminada exitosamente"}
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error al eliminar la orden: {str(e)}",
        )


# ==================== GENERACIÓN DE DOCUMENTO WORD ====================


def set_cell_border(cell, **kwargs):
    """
    Helper function to set cell borders
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        if border_name in kwargs:
            element = OxmlElement(f"w:{border_name}")
            element.set(qn("w:val"), kwargs[border_name])
            element.set(qn("w:sz"), "4")  # Size
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "auto")
            tcPr.append(element)


@router.post(
    "/{orden_id}/generar-documento", dependencies=[Depends(coordinador_or_admin)]
)
def generar_documento_orden(
    orden_id: int,
    db: Session = Depends(get_db),
):
    """
    Genera un documento Word con el reporte de la orden de trabajo.
    Incluye firmas y tabla de actividades.
    Acceso restringido a Coordinadores y Administradores.
    """
    orden = (
        db.query(OrdenTrabajoDiario).filter(OrdenTrabajoDiario.id == orden_id).first()
    )

    if not orden:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Orden de trabajo con ID {orden_id} no encontrada",
        )

    doc = Document()

    # Configurar márgenes
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Formato de fecha
    fecha_str = orden.fecha.strftime("%d/%m/%Y")

    # ==================== LOGO ====================
    import os
    logo_path = "/home/dark/Escritorio/void/mtto/backend/assets/logo.jpg"
    if os.path.exists(logo_path):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run()
        try:
            # El ancho de la página es 8.5, menos 1 de margen por cada lado = 6.5 pulgadas útiles
            logo_run.add_picture(logo_path, width=Inches(6.5))
        except Exception:
            pass

    doc.add_paragraph()  # Espacio

    # ==================== FECHA ====================
    p_fecha = doc.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_fecha = p_fecha.add_run(f"Fecha: {fecha_str}")
    run_fecha.font.size = Pt(12)

    doc.add_paragraph()  # Espacio

    # ==================== PÁRRAFO PRINCIPAL ====================
    p_main = doc.add_paragraph()
    texto_main = (
        "Yo Fernelis Barragan _________________________ le hago llegar a la oficina de talento humano encargado por la sr/a joyce pedroso.\n"
        "Donde se hace constar que se realizaron las siguientes actividades,"
    )
    run_main = p_main.add_run(texto_main)
    run_main.font.size = Pt(12)
    p_main.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()  # Espacio

    # 4. Tabla de trabajos
    # Crear lista de trabajos que tienen contenido
    trabajos = []
    if orden.trabajo_1_a_realizar or orden.trabajo_1_realizado:
        trabajos.append((orden.trabajo_1_a_realizar, orden.trabajo_1_realizado))
    if orden.trabajo_2_a_realizar or orden.trabajo_2_realizado:
        trabajos.append((orden.trabajo_2_a_realizar, orden.trabajo_2_realizado))
    if orden.trabajo_3_a_realizar or orden.trabajo_3_realizado:
        trabajos.append((orden.trabajo_3_a_realizar, orden.trabajo_3_realizado))
    if orden.trabajo_4_a_realizar or orden.trabajo_4_realizado:
        trabajos.append((orden.trabajo_4_a_realizar, orden.trabajo_4_realizado))
    if orden.trabajo_5_a_realizar or orden.trabajo_5_realizado:
        trabajos.append((orden.trabajo_5_a_realizar, orden.trabajo_5_realizado))

    if trabajos:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        # Cabecera
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "TRABAJO A REALIZAR"
        hdr_cells[1].text = "LABOR REALIZADO"

        # Estilo cabecera
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Llenar datos
        for a_realizar, realizado in trabajos:
            row_cells = table.add_row().cells
            row_cells[0].text = a_realizar if a_realizar else "N/A"
            row_cells[1].text = realizado if realizado else "N/A"

    # ==================== FIRMA FINAL ====================
    doc.add_paragraph()  # Espacio
    
    p_linea_final = doc.add_paragraph()
    p_linea_final.paragraph_format.space_before = Pt(40)
    run_linea_final = p_linea_final.add_run("______________________________")
    p_linea_final.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_firma_final = doc.add_paragraph()
    run_firma_final = p_firma_final.add_run("Fernelis Barragan")
    run_firma_final.bold = True
    p_firma_final.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Guardar en memoria
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    filename = f"Orden_Trabajo_{orden.fecha.strftime('%Y-%m-%d')}.docx"

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
