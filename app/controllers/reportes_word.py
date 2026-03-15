from fastapi import (
    APIRouter,
    Depends,
    HTTPException,
    Query,
    status,
    Form,
    File,
    UploadFile,
)
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import desc
from typing import Optional
from datetime import date
import io
import os
import urllib.request
import cloudinary
import cloudinary.uploader
from dotenv import load_dotenv

load_dotenv()

cloudinary.config(
    cloud_name=os.getenv("CLOUDINARY_CLOUD_NAME"),
    api_key=os.getenv("CLOUDINARY_API_KEY"),
    api_secret=os.getenv("CLOUDINARY_API_SECRET"),
    secure=True
)

from app.core.database import get_db
from app.core.auth_deps import coordinador_or_admin, any_auth
from app.models.models_trabajos_diarios import TrabajoDiario
from app.schemas.schemas_trabajos_diarios import (
    TrabajoDiarioSchema,
    TrabajoDiarioCreate,
    TrabajoDiarioUpdate,
    TrabajoDiarioListResponse,
)

# Importaciones para generación de Word
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


router = APIRouter(
    prefix="/trabajos-diarios",
    tags=["Trabajos Diarios"],
    responses={404: {"description": "Trabajo no encontrado"}},
)


# ==================== RUTAS GET ====================


@router.get(
    "/", response_model=TrabajoDiarioListResponse, dependencies=[Depends(any_auth)]
)
def get_trabajos_diarios(
    skip: int = Query(0, ge=0, description="Número de registros a omitir"),
    limit: int = Query(
        100, ge=1, le=1000, description="Número de registros a retornar"
    ),
    fecha_inicio: Optional[date] = Query(None, description="Filtrar desde fecha"),
    fecha_fin: Optional[date] = Query(None, description="Filtrar hasta fecha"),
    sitio: Optional[str] = Query(None, description="Filtrar por sitio de trabajo"),
    db: Session = Depends(get_db),
):
    """
    Obtiene la lista de trabajos diarios con paginación y filtros opcionales.

    - **skip**: Número de registros a omitir (para paginación)
    - **limit**: Número máximo de registros a retornar
    - **fecha_inicio**: Filtrar trabajos desde esta fecha
    - **fecha_fin**: Filtrar trabajos hasta esta fecha
    - **sitio**: Filtrar por sitio de trabajo (coincidencia parcial)
    """
    query = db.query(TrabajoDiario)

    # Aplicar filtros
    if fecha_inicio:
        query = query.filter(TrabajoDiario.fecha >= fecha_inicio)

    if fecha_fin:
        query = query.filter(TrabajoDiario.fecha <= fecha_fin)

    if sitio:
        query = query.filter(TrabajoDiario.sitio_trabajo.ilike(f"%{sitio}%"))

    # Ordenar por fecha descendente (más recientes primero)
    query = query.order_by(desc(TrabajoDiario.fecha))

    # Contar total
    total = query.count()

    # Aplicar paginación
    trabajos = query.offset(skip).limit(limit).all()

    return TrabajoDiarioListResponse(
        success=True,
        total=total,
        skip=skip,
        limit=limit,
        data=trabajos,
    )


@router.get(
    "/{trabajo_id}",
    response_model=TrabajoDiarioSchema,
    dependencies=[Depends(any_auth)],
)
def get_trabajo_by_id(
    trabajo_id: int,
    db: Session = Depends(get_db),
):
    """
    Obtiene un trabajo diario específico por su ID.

    - **trabajo_id**: ID único del trabajo diario
    """
    trabajo = db.query(TrabajoDiario).filter(TrabajoDiario.id == trabajo_id).first()

    if not trabajo:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Trabajo con ID {trabajo_id} no encontrado",
        )

    return trabajo


# ==================== RUTAS POST ====================


@router.post(
    "/",
    response_model=TrabajoDiarioSchema,
    status_code=status.HTTP_201_CREATED,
    dependencies=[Depends(any_auth)],
)
def crear_trabajo_diario(
    fecha: date = Form(...),
    sitio_trabajo: str = Form(...),
    maquinaria_trabajada: str = Form(""),
    trabajo_realizado: str = Form(...),
    fotos: list[UploadFile] = File(None),
    db: Session = Depends(get_db),
):
    """
    Crea un nuevo registro de trabajo diario.

    Validaciones:
    - La fecha, sitio y trabajo realizado son obligatorios (Datos de formulario)
    - Se pueden adjuntar hasta 8 fotografías (archivos reales)
    """

    # 1. Procesar imágenes
    rutas_fotos = []

    if fotos:
        for i, foto in enumerate(fotos):
            if i >= 8:
                break  # Limite de 8 fotos

            try:
                # Leer el contenido de la imagen
                content = foto.file.read()
                
                # Subir directamente a Cloudinary
                response = cloudinary.uploader.upload(
                    content,
                    folder="trabajos_diarios"
                )
                
                # Guardar la URL segura
                rutas_fotos.append(response.get("secure_url"))

            except Exception as e:
                print(f"Error subiendo foto {foto.filename} a Cloudinary: {e}")
                # Continuar con las siguientes fotos si una falla

    # Completar lista hasta 8 elementos con None
    while len(rutas_fotos) < 8:
        rutas_fotos.append(None)

    # 2. Crear objeto en BD
    nuevo_trabajo = TrabajoDiario(
        fecha=fecha,
        sitio_trabajo=sitio_trabajo,
        maquinaria_trabajada=maquinaria_trabajada,
        trabajo_realizado=trabajo_realizado,
        foto_1=rutas_fotos[0],
        foto_2=rutas_fotos[1],
        foto_3=rutas_fotos[2],
        foto_4=rutas_fotos[3],
        foto_5=rutas_fotos[4],
        foto_6=rutas_fotos[5],
        foto_7=rutas_fotos[6],
        foto_8=rutas_fotos[7],
    )

    try:
        db.add(nuevo_trabajo)
        db.commit()
        db.refresh(nuevo_trabajo)
        return nuevo_trabajo
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error al crear el trabajo diario: {str(e)}",
        )


# ==================== RUTAS PUT ====================


@router.put(
    "/{trabajo_id}",
    response_model=TrabajoDiarioSchema,
    dependencies=[Depends(coordinador_or_admin)],
)
def actualizar_trabajo_completo(
    trabajo_id: int,
    trabajo_data: TrabajoDiarioUpdate,
    db: Session = Depends(get_db),
):
    """
    Actualiza un trabajo diario existente (PUT).

    - **trabajo_id**: ID del trabajo a actualizar
    """
    trabajo = db.query(TrabajoDiario).filter(TrabajoDiario.id == trabajo_id).first()

    if not trabajo:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Trabajo con ID {trabajo_id} no encontrado",
        )

    # Actualizar solo los campos proporcionados
    update_data = trabajo_data.model_dump(exclude_unset=True)

    for key, value in update_data.items():
        setattr(trabajo, key, value)

    try:
        db.commit()
        db.refresh(trabajo)
        return trabajo
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error al actualizar el trabajo: {str(e)}",
        )


# ==================== RUTAS DELETE ====================


@router.delete(
    "/{trabajo_id}",
    status_code=status.HTTP_200_OK,
    dependencies=[Depends(coordinador_or_admin)],
)
def eliminar_trabajo(
    trabajo_id: int,
    db: Session = Depends(get_db),
):
    """
    Elimina un trabajo diario del sistema.

    ADVERTENCIA: Esta operación es irreversible.

    - **trabajo_id**: ID del trabajo a eliminar
    """
    trabajo = db.query(TrabajoDiario).filter(TrabajoDiario.id == trabajo_id).first()

    if not trabajo:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Trabajo con ID {trabajo_id} no encontrado",
        )

    try:
        db.delete(trabajo)
        db.commit()
        return {
            "success": True,
            "message": f"Trabajo ID {trabajo_id} eliminado exitosamente",
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error al eliminar el trabajo: {str(e)}",
        )


# ==================== GENERACIÓN DE REPORTE EN WORD ====================


@router.post("/reportes/word/{trabajo_id}")
def generar_reporte_word(
    trabajo_id: int,
    db: Session = Depends(get_db),
):
    """
    Genera un reporte en formato Word (.docx) del trabajo diario especificado.

    El reporte incluye:
    - Cabecera: "REPÚBLICA BOLIVARIANA DE VENEZUELA - ALIMENTOS Y GRASAS OCCIDENTE"
    - Logo de la empresa (si existe)
    - Fecha, sitio, maquinaria y trabajo realizado
    - Hasta 8 fotografías adjuntas (400x400 px)

    El archivo se descarga con el nombre: {fecha}_{trabajo_realizado}.docx
    """
    # Obtener el trabajo de la base de datos
    trabajo = db.query(TrabajoDiario).filter(TrabajoDiario.id == trabajo_id).first()

    if not trabajo:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Trabajo con ID {trabajo_id} no encontrado",
        )

    # Crear documento Word
    doc = Document()

    # ==================== CONFIGURAR SECCIONES ====================
    section = doc.sections[0]
    section.page_height = Inches(11)  # Carta
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # ==================== LOGO ====================
    # Intentar agregar logo si existe
    logo_path = "/home/dark/Escritorio/void/mtto/backend/assets/logo.jpg"
    if os.path.exists(logo_path):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run()
        try:
            # El ancho de la página es 8.5, menos 0.75 de margen por cada lado = 7 pulgadas útiles
            logo_run.add_picture(logo_path, width=Inches(7.0))
        except Exception:
            # Si hay error al cargar la imagen, continuar sin ella
            pass

    doc.add_paragraph()  # Espacio

    # ==================== TÍTULO DEL REPORTE ====================
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo.add_run("REPORTE DE TRABAJO DIARIO")
    titulo_run.bold = True
    titulo_run.font.size = Pt(18)
    titulo_run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()  # Espacio

    # ==================== INFORMACIÓN DEL TRABAJO ====================
    # Fecha
    fecha_para = doc.add_paragraph()
    fecha_run = fecha_para.add_run("Fecha: ")
    fecha_run.bold = True
    fecha_run.font.size = Pt(12)
    fecha_para.add_run(trabajo.fecha.strftime("%d/%m/%Y")).font.size = Pt(12)

    # Sitio
    sitio_para = doc.add_paragraph()
    sitio_run = sitio_para.add_run("Lugar: ")
    sitio_run.bold = True
    sitio_run.font.size = Pt(12)
    sitio_para.add_run(trabajo.sitio_trabajo).font.size = Pt(12)

    # Maquinaria (si existe)
    if trabajo.maquinaria_trabajada:
        maq_para = doc.add_paragraph()
        maq_run = maq_para.add_run("Maquinaria a Trabajar: ")
        maq_run.bold = True
        maq_run.font.size = Pt(12)
        maq_para.add_run(trabajo.maquinaria_trabajada).font.size = Pt(12)

    # Trabajo Realizado
    trabajo_para = doc.add_paragraph()
    trabajo_run = trabajo_para.add_run("Trabajo Realizado:\n")
    trabajo_run.bold = True
    trabajo_run.font.size = Pt(12)
    trabajo_detalle = trabajo_para.add_run(trabajo.trabajo_realizado)
    trabajo_detalle.font.size = Pt(11)

    doc.add_paragraph()  # Espacio

    # ==================== FOTOGRAFÍAS ====================
    fotos = [
        trabajo.foto_1,
        trabajo.foto_2,
        trabajo.foto_3,
        trabajo.foto_4,
        trabajo.foto_5,
        trabajo.foto_6,
        trabajo.foto_7,
        trabajo.foto_8,
    ]

    # Filtrar solo las fotos que existen (ahora son URLs de Cloudinary)
    fotos_validas = [foto for foto in fotos if foto]

    if fotos_validas:
        fotos_titulo = doc.add_paragraph()
        fotos_titulo_run = fotos_titulo.add_run("Evidencia Fotográfica:")
        fotos_titulo_run.bold = True
        fotos_titulo_run.font.size = Pt(13)

        # Agregar fotos en pares (2 por fila) para optimizar espacio
        for i in range(0, len(fotos_validas), 2):
            foto_para = doc.add_paragraph()
            foto_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Primera foto
            try:
                req = urllib.request.Request(fotos_validas[i], headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req) as response:
                    img_data = io.BytesIO(response.read())

                foto_run = foto_para.add_run()
                foto_run.add_picture(
                    img_data, width=Inches(2.8), height=Inches(2.8)
                )
            except Exception as e:
                print(f"Error cargando foto 1 para Word desede Cloudinary: {e}")
                # Si falla la carga de imagen, continuar
                pass

            # Segunda foto (si existe)
            if i + 1 < len(fotos_validas):
                try:
                    req = urllib.request.Request(fotos_validas[i+1], headers={'User-Agent': 'Mozilla/5.0'})
                    with urllib.request.urlopen(req) as response:
                        img_data = io.BytesIO(response.read())

                    foto_para.add_run("   ")  # Espacio entre fotos
                    foto_run2 = foto_para.add_run()
                    foto_run2.add_picture(
                        img_data, width=Inches(2.8), height=Inches(2.8)
                    )
                except Exception as e:
                    print(f"Error cargando foto 2 para Word desede Cloudinary: {e}")
                    pass

            doc.add_paragraph()  # Espacio entre filas

    # ==================== PIE DE PÁGINA ====================
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "___________________________\nFERNELY BARRAGAN"
    )
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)  # Gris

    # ==================== GUARDAR DOCUMENTO EN MEMORIA ====================
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # Crear nombre de archivo (sanitizar para evitar caracteres ilegales en headers)
    fecha_str = trabajo.fecha.strftime("%Y-%m-%d")
    # Remover saltos de línea, tabulaciones y caracteres especiales
    trabajo_limpio = (
        trabajo.trabajo_realizado.replace("\n", " ")
        .replace("\r", " ")
        .replace("\t", " ")
        .replace("/", "-")
        .replace("\\", "-")
        .replace(":", "-")
    )
    # Tomar solo las primeras palabras y limitar longitud
    trabajo_nombre = "_".join(trabajo_limpio.split()[:5])[:40]
    filename = f"{fecha_str}_{trabajo_nombre}.docx"

    # Retornar como descarga
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
